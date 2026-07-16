from base64 import b64decode
from io import BytesIO
from pathlib import Path
import re
import unittest
from zipfile import ZipFile

import pandas as pd
import plotly.graph_objects as go
from docx import Document
from docx.shared import Mm
from docxtpl import DocxTemplate

try:
    import report_logic
except ImportError:
    report_logic = None


EXPECTED_TEMPLATE_PLACEHOLDERS = {
    "school_name",
    "exam_name",
    "class_name",
    "subject_name",
    "generated_at",
    "student_count",
    "average_score",
    "highest_score",
    "lowest_score",
    "pass_rate",
    "excellent_rate",
    "summary_text",
    "distribution_analysis",
    "level_analysis",
    "distribution_chart",
    "level_chart",
    "subject_average_chart",
    "subject_average_fallback",
    "excellent_total_count",
    "excellent_display_students",
    "excellent_omitted_count",
    "struggling_total_count",
    "struggling_display_students",
    "struggling_omitted_count",
    "teaching_suggestions",
    "data_note",
}

# 1×1 PNG：模板契约测试只验证图片替换和 DOCX 结构，不依赖 Kaleido/Chromium。
TINY_PNG = b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
)


class ReportLogicTest(unittest.TestCase):
    def setUp(self):
        self.analysis_result = {
            "student_count": 3,
            "average_score": 88.5,
            "highest_score": 100.0,
            "lowest_score": 72.0,
            "pass_rate": 66.6666,
            "excellent_rate": 33.3333,
        }
        self.excellent_df = pd.DataFrame([["张三&甲", 100]], columns=["姓名", "分数"])
        self.fail_df = pd.DataFrame([["李<四", 58]], columns=["姓名", "分数"])
        self.distribution_figure = go.Figure(go.Bar(x=["及格", "优秀"], y=[2, 1]))
        self.level_figure = go.Figure(go.Pie(labels=["及格", "优秀"], values=[2, 1]))
        self.distribution = pd.DataFrame(
            {
                "档位": ["待提升", "及格", "中等", "良好", "优秀"],
                "区间": ["[0, 72)", "[72, 84)", "[84, 96)", "[96, 108)", "[108, 120]"],
                "人数": [1, 0, 1, 0, 1],
                "占比": [33.3333, 0.0, 33.3333, 0.0, 33.3333],
            }
        )
        self.distribution.attrs["有效人数"] = 3

    def require_report_logic(self):
        self.assertIsNotNone(report_logic, "缺少 report_logic 模块")

    @staticmethod
    def _all_paragraphs(document):
        yield from document.paragraphs
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from cell.paragraphs
        for section in document.sections:
            yield from section.header.paragraphs
            yield from section.first_page_header.paragraphs
            yield from section.footer.paragraphs
            yield from section.first_page_footer.paragraphs

    def test_template_uses_explicit_page_breaks_core_metric_table_and_atomic_placeholders(self):
        self.require_report_logic()

        template = Document(report_logic.DEFAULT_TEMPLATE_PATH)
        template_contract = DocxTemplate(report_logic.DEFAULT_TEMPLATE_PATH)
        page_break_count = len(
            template.element.body.xpath('.//w:br[@w:type="page"]')
        )
        table_shapes = {(len(table.rows), len(table.columns)) for table in template.tables}
        jinja_expressions = [
            match.group(0)
            for paragraph in self._all_paragraphs(template)
            for match in re.finditer(r"\{\{.*?\}\}|\{%.*?%\}", paragraph.text)
        ]

        self.assertEqual(page_break_count, 4)
        self.assertIn((4, 3), table_shapes)
        self.assertIn((4, 4), table_shapes)
        self.assertEqual(
            set(template_contract.get_undeclared_template_variables()),
            EXPECTED_TEMPLATE_PLACEHOLDERS,
        )
        for expression in jinja_expressions:
            self.assertTrue(
                any(
                    run.text.strip() == expression.strip()
                    for paragraph in self._all_paragraphs(template)
                    for run in paragraph.runs
                ),
                f"Jinja 表达式必须位于独立完整 Word run：{expression}",
            )
        self.assertEqual(
            set(report_logic.TEMPLATE_CONTEXT_VARIABLES),
            EXPECTED_TEMPLATE_PLACEHOLDERS,
        )

    def test_context_uses_generated_at_and_preserves_specific_or_all_student_scope(self):
        self.require_report_logic()

        specific_context = report_logic.build_report_context(
            analysis_result=self.analysis_result,
            excellent_df=self.excellent_df,
            fail_df=self.fail_df,
            distribution=self.distribution,
            selected_class="2501",
            score_col="数学",
            full_score=120,
            school_name="示例学校",
            exam_name="期中考试",
            generated_at="2026年7月17日 00:35",
        )
        all_students_context = report_logic.build_report_context(
            analysis_result=self.analysis_result,
            excellent_df=self.excellent_df,
            fail_df=self.fail_df,
            distribution=self.distribution,
            selected_class="全部学生",
            score_col="语文",
            full_score=120,
            school_name="示例学校",
            exam_name="期中考试",
            generated_at="2026年7月17日 00:35",
        )

        self.assertEqual(specific_context["class_name"], "2501")
        self.assertEqual(all_students_context["class_name"], "全部学生")
        self.assertEqual(specific_context["generated_at"], "2026年7月17日 00:35")
        self.assertEqual(all_students_context["generated_at"], "2026年7月17日 00:35")

    def test_student_lists_are_sorted_limited_and_do_not_mutate_page_data(self):
        self.require_report_logic()
        excellent_df = pd.DataFrame(
            [[f"优{i:02d}", 80 + i] for i in range(12)],
            columns=["姓名", "分数"],
        )
        fail_df = pd.DataFrame(
            [[f"待{i:02d}", 59 - i] for i in range(13)],
            columns=["姓名", "分数"],
        )
        original_excellent = excellent_df.copy(deep=True)
        original_fail = fail_df.copy(deep=True)

        context = report_logic.build_report_context(
            analysis_result=self.analysis_result,
            excellent_df=excellent_df,
            fail_df=fail_df,
            distribution=self.distribution,
            selected_class="全部学生",
            score_col="数学",
            full_score=120,
            school_name="示例学校",
            exam_name="期中考试",
            generated_at="2026年7月17日 00:35",
        )

        self.assertEqual(context["excellent_total_count"], 12)
        self.assertEqual(context["excellent_omitted_count"], 2)
        self.assertEqual(context["struggling_total_count"], 13)
        self.assertEqual(context["struggling_omitted_count"], 3)
        self.assertEqual(
            [row["score"] for row in context["excellent_display_students"] if row["name"]],
            ["91", "90", "89", "88", "87", "86", "85", "84", "83", "82"],
        )
        self.assertEqual(
            [row["score"] for row in context["struggling_display_students"] if row["name"]],
            ["47", "48", "49", "50", "51", "52", "53", "54", "55", "56"],
        )
        pd.testing.assert_frame_equal(excellent_df, original_excellent)
        pd.testing.assert_frame_equal(fail_df, original_fail)

    def test_context_data_note_uses_current_scoring_settings(self):
        self.require_report_logic()
        analysis_result = dict(self.analysis_result, excellent_percent=85)

        context = report_logic.build_report_context(
            analysis_result=analysis_result,
            excellent_df=self.excellent_df,
            fail_df=self.fail_df,
            distribution=self.distribution,
            selected_class="2501",
            score_col="语文",
            full_score=120,
            school_name="示例学校",
            exam_name="期中考试",
            generated_at="2026年7月17日 00:35",
        )

        self.assertIn("当前分析范围为 2501", context["data_note"])
        self.assertIn("当前科目或成绩列为语文", context["data_note"])
        self.assertIn("当前成绩列满分为 120 分", context["data_note"])
        self.assertIn("及格线为满分的 60%", context["data_note"])
        self.assertIn("优秀线为满分的 85%", context["data_note"])
        self.assertIn("完整名单请查看 Excel 导出", context["data_note"])

    def test_template_contract_can_render_with_simulated_data(self):
        """手动修改模板后可单独运行的轻量完整性检查。"""
        self.require_report_logic()

        template = Document(report_logic.DEFAULT_TEMPLATE_PATH)
        template_contract = DocxTemplate(report_logic.DEFAULT_TEMPLATE_PATH)
        template_paragraphs = list(self._all_paragraphs(template))
        self.assertEqual(
            set(template_contract.get_undeclared_template_variables()),
            EXPECTED_TEMPLATE_PLACEHOLDERS,
        )
        for paragraph in template_paragraphs:
            for match in re.finditer(r"\{\{.*?\}\}|\{%.*?%\}", paragraph.text):
                self.assertTrue(
                    any(run.text.strip() == match.group(0).strip() for run in paragraph.runs),
                    f"Jinja 表达式必须位于独立完整 Word run：{match.group(0)}",
                )

        class StaticPngFigure:
            def __init__(self, marker):
                self.marker = marker

            def write_image(self, output_path, **_kwargs):
                # PNG 允许尾随数据；不同标记可避免 python-docx 合并两张相同图片。
                Path(output_path).write_bytes(TINY_PNG + self.marker)

        report_bytes = report_logic.build_score_report_bytes(
            analysis_result=self.analysis_result,
            excellent_df=self.excellent_df,
            fail_df=self.fail_df,
            distribution=self.distribution,
            distribution_figure=StaticPngFigure(b"distribution"),
            level_figure=StaticPngFigure(b"level"),
            selected_class="模板验证班",
            score_col="数学",
            full_score=120,
            school_name="模板验证学校",
            exam_name="模板验证考试",
            generated_at="2026年7月16日 09:30",
        )

        with ZipFile(BytesIO(report_bytes)) as archive:
            self.assertIsNone(archive.testzip())
            media_names = [
                name for name in archive.namelist() if name.startswith("word/media/")
            ]
            word_xml_parts = [
                archive.read(name)
                for name in archive.namelist()
                if name.startswith("word/") and name.endswith(".xml")
            ]

        rendered_document = Document(BytesIO(report_bytes))
        rendered_text = "\n".join(
            paragraph.text for paragraph in self._all_paragraphs(rendered_document)
        )
        self.assertEqual(len(media_names), 2)
        self.assertTrue(
            all(b"{{" not in xml_part and b"}}" not in xml_part for xml_part in word_xml_parts)
        )
        self.assertIn("模板验证学校", rendered_text)
        self.assertIn("模板验证考试", rendered_text)

    def test_generated_report_keeps_images_bounded_and_summarizes_40_names(self):
        self.require_report_logic()

        analysis_result = {
            "student_count": 40,
            "average_score": 78.45,
            "highest_score": 118.0,
            "lowest_score": 32.0,
            "pass_rate": 72.5,
            "excellent_rate": 22.5,
            "fail_count": 11,
            "excellent_count": 9,
        }
        excellent_names = [f"优秀学生{i:02d}" for i in range(1, 21)]
        struggling_names = [f"待提升学生{i:02d}" for i in range(21, 41)]
        excellent_df = pd.DataFrame(
            [[name, 120 - index] for index, name in enumerate(excellent_names)],
            columns=["姓名", "分数"],
        )
        fail_df = pd.DataFrame(
            [[name, 59 - index] for index, name in enumerate(struggling_names)],
            columns=["姓名", "分数"],
        )

        report_bytes = report_logic.build_score_report_bytes(
            analysis_result=analysis_result,
            excellent_df=excellent_df,
            fail_df=fail_df,
            distribution=self.distribution,
            distribution_figure=self.distribution_figure,
            level_figure=self.level_figure,
            selected_class="初二（3）班",
            score_col="语文",
            full_score=120,
            school_name="示例中学",
            exam_name="期中考试",
            generated_at="2026年7月15日 09:30",
        )

        document = Document(BytesIO(report_bytes))
        section = document.sections[0]
        content_width = section.page_width - section.left_margin - section.right_margin
        rendered_text = "\n".join(
            paragraph.text for paragraph in self._all_paragraphs(document)
        )

        self.assertEqual(len(document.inline_shapes), 2)
        self.assertTrue(all(shape.width <= content_width for shape in document.inline_shapes))
        self.assertTrue(all(shape.width <= Mm(155) for shape in document.inline_shapes))
        self.assertEqual(
            len(document.element.body.xpath('.//w:br[@w:type="page"]')),
            4,
        )
        for name in excellent_names[:10] + list(reversed(struggling_names))[:10]:
            self.assertIn(name, rendered_text)
        for name in excellent_names[10:] + struggling_names[:10]:
            self.assertNotIn(name, rendered_text)
        self.assertIn("优秀学生共20人。", rendered_text)
        self.assertIn("待提升学生共20人。", rendered_text)
        self.assertEqual(
            rendered_text.count("本报告展示前10人，另有10人未列出"),
            2,
        )

    def test_report_without_subject_average_figure_uses_fallback_and_two_images(self):
        self.require_report_logic()

        report_bytes = report_logic.build_score_report_bytes(
            analysis_result=self.analysis_result,
            excellent_df=self.excellent_df,
            fail_df=self.fail_df,
            distribution=self.distribution,
            distribution_figure=self.distribution_figure,
            level_figure=self.level_figure,
            subject_average_figure=None,
            selected_class="2501",
            score_col="数学",
            full_score=120,
            school_name="示例学校",
            exam_name="期中考试",
            generated_at="2026年7月17日 00:35",
        )

        document = Document(BytesIO(report_bytes))
        rendered_text = "\n".join(
            paragraph.text for paragraph in self._all_paragraphs(document)
        )
        self.assertEqual(len(document.inline_shapes), 2)
        self.assertIn(
            "当前数据仅识别到一个有效成绩科目，暂无各科平均分对比。",
            rendered_text,
        )

    def test_report_with_subject_average_figure_embeds_three_bounded_images(self):
        self.require_report_logic()
        subject_average_figure = go.Figure(
            go.Bar(x=["语文", "数学", "英语"], y=[88, 90, 86])
        )

        report_bytes = report_logic.build_score_report_bytes(
            analysis_result=self.analysis_result,
            excellent_df=self.excellent_df,
            fail_df=self.fail_df,
            distribution=self.distribution,
            distribution_figure=self.distribution_figure,
            level_figure=self.level_figure,
            subject_average_figure=subject_average_figure,
            selected_class="全部学生",
            score_col="数学",
            full_score=120,
            school_name="示例学校",
            exam_name="期中考试",
            generated_at="2026年7月17日 00:35",
        )

        document = Document(BytesIO(report_bytes))
        self.assertEqual(len(document.inline_shapes), 3)
        self.assertTrue(all(shape.width <= Mm(155) for shape in document.inline_shapes))

    def test_app_passes_optional_subject_average_figure_to_word_report(self):
        source = Path("app.py").read_text(encoding="utf-8")

        self.assertIn("subject_average_figure = None", source)
        self.assertIn("subject_average_figure=subject_average_figure", source)

    def test_build_report_context_formats_current_web_values_and_empty_inputs(self):
        self.require_report_logic()

        context = report_logic.build_report_context(
            analysis_result=self.analysis_result,
            excellent_df=self.excellent_df,
            fail_df=self.fail_df,
            distribution=self.distribution,
            selected_class="2401",
            score_col="数学分数",
            full_score=120,
            school_name="",
            exam_name="",
            generated_at="2026年7月14日 08:15",
        )

        self.assertEqual(context["school_name"], "未填写")
        self.assertEqual(context["exam_name"], "成绩分析")
        self.assertEqual(context["class_name"], "2401")
        self.assertEqual(context["subject_name"], "数学分数")
        self.assertEqual(context["student_count"], "3")
        self.assertEqual(context["average_score"], "88.50")
        self.assertEqual(context["highest_score"], "100")
        self.assertEqual(context["lowest_score"], "72")
        self.assertEqual(context["pass_rate"], "66.7%")
        self.assertEqual(context["excellent_rate"], "33.3%")
        self.assertEqual(context["generated_at"], "2026年7月14日 08:15")
        self.assertEqual(context["excellent_display_students"][0]["name"], "张三&甲")
        self.assertEqual(context["excellent_display_students"][0]["score"], "100")
        self.assertEqual(context["struggling_display_students"][0]["name"], "李<四")
        self.assertEqual(context["struggling_display_students"][0]["score"], "58")
        self.assertIn("平均得分率为 73.8%", context["summary_text"])
        self.assertIn("极差为 28 分", context["summary_text"])
        self.assertIn("样本较少，结论仅供参考", context["summary_text"])
        self.assertNotIn("总分分析反映整体结果", context["summary_text"])
        self.assertNotIn("成绩较差", context["summary_text"] + context["teaching_suggestions"])

    def test_build_report_context_uses_empty_list_fallbacks(self):
        self.require_report_logic()

        context = report_logic.build_report_context(
            analysis_result=self.analysis_result,
            excellent_df=pd.DataFrame(columns=["姓名", "分数"]),
            fail_df=pd.DataFrame(columns=["姓名", "分数"]),
            distribution=self.distribution,
            selected_class="全部学生",
            score_col="总成绩分数",
            full_score=800,
            school_name="学校",
            exam_name="考试",
            generated_at="2026年7月14日 08:15",
        )

        self.assertEqual(context["excellent_total_count"], 0)
        self.assertEqual(context["struggling_total_count"], 0)
        self.assertEqual(context["excellent_omitted_count"], 0)
        self.assertEqual(context["struggling_omitted_count"], 0)
        self.assertEqual(context["excellent_display_students"][0]["name"], "暂无优秀学生")
        self.assertEqual(context["struggling_display_students"][0]["name"], "暂无待提升学生")
        self.assertIn(
            "总分分析反映整体结果，具体薄弱学科仍需结合单科数据进一步判断。",
            context["summary_text"],
        )
        self.assertIn("结合单科数据进一步判断", context["teaching_suggestions"])

    def test_distribution_analysis_identifies_tied_dominant_intervals_and_skips_invalid_rows(self):
        self.require_report_logic()
        distribution = pd.DataFrame(
            {
                "档位": ["待提升", "及格", "中等", "良好", "优秀", "无效"],
                "区间": ["0–59", "60–69", "70–79", "80–89", "90–100", "无区间"],
                "人数": [2, 8, 8, 1, 1, 0],
                "占比": [10.0, 40.0, 40.0, 5.0, 5.0, 0.0],
            }
        )

        text = report_logic._build_distribution_analysis(distribution)

        self.assertIn("主要集中在60–69、70–79两个区间", text)
        self.assertNotIn("无区间", text)
        self.assertNotIn("两极分化", text)

    def test_distribution_analysis_describes_majority_concentration(self):
        self.require_report_logic()
        distribution = self.distribution.copy()
        distribution["人数"] = [1, 1, 7, 1, 0]
        distribution["占比"] = [10.0, 10.0, 70.0, 10.0, 0.0]

        text = report_logic._build_distribution_analysis(distribution)

        self.assertIn("成绩明显集中于[84, 96)区间", text)
        self.assertNotIn("两极分化", text)

    def test_distribution_analysis_uses_strict_polarization_rule(self):
        self.require_report_logic()
        distribution = self.distribution.copy()
        distribution["人数"] = [30, 15, 10, 15, 30]
        distribution["占比"] = [30.0, 15.0, 10.0, 15.0, 30.0]

        polarized = report_logic._build_distribution_analysis(distribution)

        distribution["人数"] = [25, 30, 10, 10, 25]
        distribution["占比"] = [25.0, 30.0, 10.0, 10.0, 25.0]
        not_polarized = report_logic._build_distribution_analysis(distribution)

        self.assertIn("两极分化", polarized)
        self.assertNotIn("两极分化", not_polarized)
        self.assertIn("分布较分散", not_polarized)

    def test_level_analysis_reports_all_levels_ties_and_zero_counts_naturally(self):
        self.require_report_logic()
        distribution = self.distribution.copy()
        distribution["人数"] = [0, 4, 4, 2, 0]
        distribution["占比"] = [0.0, 40.0, 40.0, 20.0, 0.0]

        text = report_logic._build_level_analysis(distribution)

        self.assertIn("待提升层暂未形成（0人，0.0%）", text)
        self.assertIn("优秀层暂未形成（0人，0.0%）", text)
        self.assertIn("主体等级为及格和中等", text)
        for level in ("待提升", "及格", "中等", "良好", "优秀"):
            self.assertIn(level, text)

    def test_level_analysis_describes_three_tied_levels_naturally(self):
        self.require_report_logic()
        distribution = self.distribution.copy()
        distribution["人数"] = [0, 3, 3, 3, 0]
        distribution["占比"] = [0.0, 33.3333, 33.3333, 33.3333, 0.0]

        text = report_logic._build_level_analysis(distribution)

        self.assertIn("主体等级为及格、中等和良好，三个层级人数并列最多", text)

    def test_safe_report_filename_removes_windows_illegal_characters(self):
        self.require_report_logic()

        filename = report_logic.safe_report_filename(
            school_name='学/校:*?',
            class_name='2401|班',
            subject_name='数学<分数>',
            exam_name='期"末\\考',
        )

        self.assertEqual(filename, "学校_2401班_数学分数_期末考_成绩分析报告.docx")

    def test_missing_template_raises_friendly_error(self):
        self.require_report_logic()

        with self.assertRaisesRegex(report_logic.ReportGenerationError, "Word 报告模板不存在"):
            report_logic.build_score_report_bytes(
                analysis_result=self.analysis_result,
                excellent_df=self.excellent_df,
                fail_df=self.fail_df,
                distribution=self.distribution,
                distribution_figure=self.distribution_figure,
                level_figure=self.level_figure,
                selected_class="2401",
                score_col="数学",
                full_score=120,
                school_name="学校",
                exam_name="考试",
                template_path=Path("missing-template.docx"),
            )

    def test_chart_export_failure_does_not_return_partial_report(self):
        self.require_report_logic()

        class FailingFigure:
            def write_image(self, *_args, **_kwargs):
                raise RuntimeError("browser unavailable")

        with self.assertRaisesRegex(report_logic.ReportGenerationError, "Word 报告图表生成失败"):
            report_logic.build_score_report_bytes(
                analysis_result=self.analysis_result,
                excellent_df=self.excellent_df,
                fail_df=self.fail_df,
                distribution=self.distribution,
                distribution_figure=FailingFigure(),
                level_figure=FailingFigure(),
                selected_class="2401",
                score_col="数学",
                full_score=120,
                school_name="学校",
                exam_name="考试",
            )

    def test_generated_report_is_valid_docx_with_two_images_and_no_placeholders(self):
        self.require_report_logic()

        report_bytes = report_logic.build_score_report_bytes(
            analysis_result=self.analysis_result,
            excellent_df=self.excellent_df,
            fail_df=self.fail_df,
            distribution=self.distribution,
            distribution_figure=self.distribution_figure,
            level_figure=self.level_figure,
            selected_class="2401",
            score_col="数学",
            full_score=120,
            school_name="示例学校",
            exam_name="期末考试",
            report_date="2026-07-14",
        )

        with ZipFile(BytesIO(report_bytes)) as archive:
            media_names = [name for name in archive.namelist() if name.startswith("word/media/")]
            document_xml = archive.read("word/document.xml").decode("utf-8")
        document = Document(BytesIO(report_bytes))
        rendered_text = "\n".join(
            paragraph.text for paragraph in self._all_paragraphs(document)
        )

        self.assertEqual(len(media_names), 2)
        self.assertNotIn("{{", document_xml)
        self.assertNotIn("}}", document_xml)
        self.assertIn("示例学校", rendered_text)
        self.assertIn("张三&甲", rendered_text)
        self.assertIn("李<四", rendered_text)

    def test_report_generation_leaves_no_project_png_or_docx(self):
        self.require_report_logic()
        project_root = Path(__file__).resolve().parent
        before = {path.name for path in project_root.glob("*.png")} | {
            path.name for path in project_root.glob("*.docx")
        }
        report_logic.build_score_report_bytes(
            analysis_result=self.analysis_result,
            excellent_df=self.excellent_df,
            fail_df=self.fail_df,
            distribution=self.distribution,
            distribution_figure=self.distribution_figure,
            level_figure=self.level_figure,
            selected_class="2401",
            score_col="数学",
            full_score=120,
            school_name="学校",
            exam_name="考试",
        )
        after = {path.name for path in project_root.glob("*.png")} | {
            path.name for path in project_root.glob("*.docx")
        }
        self.assertEqual(after, before)


if __name__ == "__main__":
    unittest.main()
